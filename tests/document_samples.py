"""In-memory document samples generated for parser and ingestion tests."""

from io import BytesIO

from docx import Document


def build_pdf_bytes(text: str = "Refund policy") -> bytes:
    """Build one minimal Type1-font PDF without storing a binary fixture."""
    encoded_text = (
        text.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)").encode("ascii")
    )
    content = b"BT /F1 12 Tf 72 720 Td (" + encoded_text + b") Tj ET"
    objects = (
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        (
            b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
            b"/Resources << /Font << /F1 5 0 R >> >> /Contents 4 0 R >>"
        ),
        b"<< /Length "
        + str(len(content)).encode("ascii")
        + b" >>\nstream\n"
        + content
        + b"\nendstream",
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
    )
    output = bytearray(b"%PDF-1.4\n%\xe2\xe3\xcf\xd3\n")
    offsets = [0]
    for object_number, body in enumerate(objects, start=1):
        offsets.append(len(output))
        output.extend(f"{object_number} 0 obj\n".encode())
        output.extend(body)
        output.extend(b"\nendobj\n")
    xref_offset = len(output)
    output.extend(f"xref\n0 {len(objects) + 1}\n".encode())
    output.extend(b"0000000000 65535 f \n")
    for offset in offsets[1:]:
        output.extend(f"{offset:010d} 00000 n \n".encode())
    output.extend(
        (
            f"trailer\n<< /Size {len(objects) + 1} /Root 1 0 R >>\n"
            f"startxref\n{xref_offset}\n%%EOF\n"
        ).encode()
    )
    return bytes(output)


def build_docx_bytes() -> bytes:
    """Build a DOCX containing heading, paragraph, list, and table structures."""
    document = Document()
    document.add_heading("Refund policy", level=1)
    document.add_paragraph("Customers may request a refund within seven days.")
    document.add_paragraph("Keep the order number.", style="List Bullet")
    table = document.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "Region"
    table.rows[0].cells[1].text = "Window"
    table.rows[1].cells[0].text = "CN"
    table.rows[1].cells[1].text = "7 days"
    output = BytesIO()
    document.save(output)
    return output.getvalue()


CSV_SAMPLE = b"name,policy\nAlice,7 days\nBob,14 days\n"
